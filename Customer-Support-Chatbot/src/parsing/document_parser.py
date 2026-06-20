"""
Document Parser - Handles multiple file formats.
Supports rule-based and AI-based parsing strategies.
"""

import hashlib
import io
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Any, Optional

import pdfplumber
import structlog
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from unstructured.partition.auto import partition
from unstructured.partition.pdf import partition_pdf

logger = structlog.get_logger(__name__)


class ParseStrategy(str, Enum):
    RULE_BASED = "rule_based"
    AI_BASED = "ai_based"       # Uses unstructured.io AI models
    HYBRID = "hybrid"


class DocumentType(str, Enum):
    PDF = "pdf"
    DOCX = "docx"
    HTML = "html"
    MARKDOWN = "markdown"
    TXT = "txt"
    CSV = "csv"
    UNKNOWN = "unknown"


@dataclass
class ParsedDocument:
    """Represents a parsed document with metadata."""
    doc_id: str
    source_path: str
    doc_type: DocumentType
    title: str
    content: str
    sections: list[dict[str, Any]] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)
    parse_strategy: str = ""
    page_count: int = 0
    
    @property
    def content_hash(self) -> str:
        return hashlib.md5(self.content.encode()).hexdigest()
    
    def to_dict(self) -> dict:
        return {
            "doc_id": self.doc_id,
            "source_path": self.source_path,
            "doc_type": self.doc_type.value,
            "title": self.title,
            "content": self.content,
            "sections": self.sections,
            "metadata": self.metadata,
            "parse_strategy": self.parse_strategy,
            "page_count": self.page_count,
            "content_hash": self.content_hash,
        }


class BaseParser(ABC):
    """Abstract base parser."""
    
    @abstractmethod
    def parse(self, file_path: Path, **kwargs) -> ParsedDocument:
        pass
    
    def _generate_doc_id(self, file_path: Path) -> str:
        return hashlib.sha256(str(file_path.resolve()).encode()).hexdigest()[:16]
    
    def _clean_text(self, text: str) -> str:
        """Basic text cleaning."""
        # Remove excessive whitespace
        import re
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        text = text.strip()
        return text


class PDFParser(BaseParser):
    """
    Rule-based PDF parser using pdfplumber.
    Extracts text with layout awareness.
    """
    
    def parse(self, file_path: Path, **kwargs) -> ParsedDocument:
        logger.info("Parsing PDF", path=str(file_path))
        
        sections = []
        full_text_parts = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text with layout
                    text = page.extract_text(
                        x_tolerance=3,
                        y_tolerance=3,
                        layout=True,
                        x_density=7.25,
                        y_density=13,
                    )
                    
                    if text:
                        cleaned = self._clean_text(text)
                        full_text_parts.append(cleaned)
                        sections.append({
                            "type": "page",
                            "page_number": page_num,
                            "content": cleaned,
                            "tables": self._extract_tables(page),
                        })
                
                # Extract metadata
                metadata = {
                    "page_count": page_count,
                    "file_size": file_path.stat().st_size,
                    "file_name": file_path.name,
                }
                if pdf.metadata:
                    metadata.update({
                        k.lower(): v 
                        for k, v in pdf.metadata.items() 
                        if v
                    })
                
        except Exception as e:
            logger.error("PDF parsing failed", path=str(file_path), error=str(e))
            raise
        
        full_content = "\n\n".join(full_text_parts)
        title = metadata.get("title", file_path.stem)
        
        return ParsedDocument(
            doc_id=self._generate_doc_id(file_path),
            source_path=str(file_path),
            doc_type=DocumentType.PDF,
            title=title,
            content=full_content,
            sections=sections,
            metadata=metadata,
            parse_strategy=ParseStrategy.RULE_BASED.value,
            page_count=page_count,
        )
    
    def _extract_tables(self, page) -> list[dict]:
        """Extract tables from a PDF page."""
        tables = []
        try:
            for table in page.extract_tables():
                if table:
                    tables.append({
                        "rows": table,
                        "row_count": len(table),
                        "col_count": len(table[0]) if table else 0,
                    })
        except Exception:
            pass
        return tables


class AIBasedPDFParser(BaseParser):
    """
    AI-based PDF parser using unstructured.io.
    Better at handling complex layouts, forms, and mixed content.
    """
    
    def parse(self, file_path: Path, **kwargs) -> ParsedDocument:
        logger.info("AI-parsing PDF", path=str(file_path))
        
        try:
            # Use unstructured's AI-based partitioning
            elements = partition_pdf(
                filename=str(file_path),
                strategy="hi_res",           # High resolution AI extraction
                infer_table_structure=True,  # AI table detection
                include_page_breaks=True,
                extract_images_in_pdf=False,
            )
            
            sections = []
            full_text_parts = []
            current_section = None
            
            for element in elements:
                elem_type = type(element).__name__
                elem_text = str(element).strip()
                
                if not elem_text:
                    continue
                
                if elem_type == "Title":
                    if current_section:
                        sections.append(current_section)
                    current_section = {
                        "type": "section",
                        "title": elem_text,
                        "content": "",
                        "elements": [],
                    }
                elif current_section:
                    current_section["content"] += f"\n{elem_text}"
                    current_section["elements"].append({
                        "type": elem_type,
                        "text": elem_text,
                    })
                
                full_text_parts.append(elem_text)
            
            if current_section:
                sections.append(current_section)
            
            full_content = "\n\n".join(full_text_parts)
            
            # Extract title from first Title element or filename
            title = (
                next((s["title"] for s in sections if s.get("title")), None)
                or file_path.stem
            )
            
            return ParsedDocument(
                doc_id=self._generate_doc_id(file_path),
                source_path=str(file_path),
                doc_type=DocumentType.PDF,
                title=title,
                content=full_content,
                sections=sections,
                metadata={
                    "file_name": file_path.name,
                    "file_size": file_path.stat().st_size,
                    "element_count": len(elements),
                },
                parse_strategy=ParseStrategy.AI_BASED.value,
            )
            
        except Exception as e:
            logger.error("AI PDF parsing failed", path=str(file_path), error=str(e))
            logger.info("Falling back to rule-based parsing")
            return PDFParser().parse(file_path, **kwargs)


class HTMLParser(BaseParser):
    """HTML document parser with structure preservation."""
    
    def parse(self, file_path: Path, **kwargs) -> ParsedDocument:
        logger.info("Parsing HTML", path=str(file_path))
        
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, "lxml")
        
        # Remove script and style elements
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        
        # Extract title
        title_tag = soup.find("title") or soup.find("h1")
        title = title_tag.get_text(strip=True) if title_tag else file_path.stem
        
        # Extract structured sections
        sections = []
        for heading in soup.find_all(["h1", "h2", "h3", "h4"]):
            section_title = heading.get_text(strip=True)
            section_content_parts = []
            
            # Gather content until next heading of same or higher level
            for sibling in heading.next_siblings:
                if hasattr(sibling, 'name') and sibling.name in ["h1", "h2", "h3", "h4"]:
                    break
                if hasattr(sibling, 'get_text'):
                    text = sibling.get_text(separator=" ", strip=True)
                    if text:
                        section_content_parts.append(text)
            
            if section_content_parts:
                sections.append({
                    "type": "section",
                    "title": section_title,
                    "heading_level": heading.name,
                    "content": " ".join(section_content_parts),
                })
        
        # Full text extraction
        full_content = soup.get_text(separator="\n", strip=True)
        full_content = self._clean_text(full_content)
        
        return ParsedDocument(
            doc_id=self._generate_doc_id(file_path),
            source_path=str(file_path),
            doc_type=DocumentType.HTML,
            title=title,
            content=full_content,
            sections=sections,
            metadata={
                "file_name": file_path.name,
                "url": kwargs.get("url", ""),
            },
            parse_strategy=ParseStrategy.RULE_BASED.value,
        )


class DocxParser(BaseParser):
    """Microsoft Word document parser."""
    
    def parse(self, file_path: Path, **kwargs) -> ParsedDocument:
        logger.info("Parsing DOCX", path=str(file_path))
        
        doc = DocxDocument(file_path)
        sections = []
        full_text_parts = []
        current_section = None
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Detect headings
            if para.style.name.startswith("Heading"):
                if current_section:
                    sections.append(current_section)
                current_section = {
                    "type": "section",
                    "title": text,
                    "content": "",
                    "style": para.style.name,
                }
            elif current_section:
                current_section["content"] += f"\n{text}"
            
            full_text_parts.append(text)
        
        if current_section:
            sections.append(current_section)
        
        # Extract tables
        table_texts = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            table_texts.append("\n".join(rows))
        
        if table_texts:
            full_text_parts.extend(table_texts)
        
        # Get title from core properties or first heading
        title = (
            doc.core_properties.title
            or (sections[0]["title"] if sections else None)
            or file_path.stem
        )
        
        return ParsedDocument(
            doc_id=self._generate_doc_id(file_path),
            source_path=str(file_path),
            doc_type=DocumentType.DOCX,
            title=title,
            content="\n\n".join(full_text_parts),
            sections=sections,
            metadata={
                "file_name": file_path.name,
                "author": doc.core_properties.author or "",
                "created": str(doc.core_properties.created or ""),
            },
            parse_strategy=ParseStrategy.RULE_BASED.value,
        )


class MarkdownParser(BaseParser):
    """Markdown document parser."""
    
    def parse(self, file_path: Path, **kwargs) -> ParsedDocument:
        logger.info("Parsing Markdown", path=str(file_path))
        import re
        
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        
        sections = []
        current_section = None
        lines = content.split("\n")
        
        for line in lines:
            # Detect headings
            heading_match = re.match(r'^(#{1,4})\s+(.+)$', line)
            if heading_match:
                if current_section:
                    sections.append(current_section)
                level = len(heading_match.group(1))
                title = heading_match.group(2).strip()
                current_section = {
                    "type": "section",
                    "title": title,
                    "level": level,
                    "content": "",
                }
            elif current_section:
                current_section["content"] += f"\n{line}"
        
        if current_section:
            sections.append(current_section)
        
        # Extract document title
        title = (
            next((s["title"] for s in sections if s.get("level") == 1), None)
            or file_path.stem
        )
        
        # Clean markdown syntax for plain text
        plain_text = re.sub(r'#{1,6}\s+', '', content)
        plain_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', plain_text)
        plain_text = re.sub(r'\*([^*]+)\*', r'\1', plain_text)
        plain_text = re.sub(r'`([^`]+)`', r'\1', plain_text)
        plain_text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', plain_text)
        
        return ParsedDocument(
            doc_id=self._generate_doc_id(file_path),
            source_path=str(file_path),
            doc_type=DocumentType.MARKDOWN,
            title=title,
            content=self._clean_text(plain_text),
            sections=sections,
            metadata={"file_name": file_path.name},
            parse_strategy=ParseStrategy.RULE_BASED.value,
        )


class DocumentParserFactory:
    """
    Factory that selects the appropriate parser based on file type and strategy.
    """
    
    _PARSERS = {
        DocumentType.PDF: {
            ParseStrategy.RULE_BASED: PDFParser,
            ParseStrategy.AI_BASED: AIBasedPDFParser,
        },
        DocumentType.HTML: {
            ParseStrategy.RULE_BASED: HTMLParser,
            ParseStrategy.AI_BASED: HTMLParser,  # HTML AI parsing fallback
        },
        DocumentType.DOCX: {
            ParseStrategy.RULE_BASED: DocxParser,
            ParseStrategy.AI_BASED: DocxParser,
        },
        DocumentType.MARKDOWN: {
            ParseStrategy.RULE_BASED: MarkdownParser,
            ParseStrategy.AI_BASED: MarkdownParser,
        },
    }
    
    _EXTENSION_MAP = {
        ".pdf": DocumentType.PDF,
        ".html": DocumentType.HTML,
        ".htm": DocumentType.HTML,
        ".docx": DocumentType.DOCX,
        ".doc": DocumentType.DOCX,
        ".md": DocumentType.MARKDOWN,
        ".markdown": DocumentType.MARKDOWN,
        ".txt": DocumentType.TXT,
    }
    
    @classmethod
    def get_doc_type(cls, file_path: Path) -> DocumentType:
        return cls._EXTENSION_MAP.get(
            file_path.suffix.lower(), DocumentType.UNKNOWN
        )
    
    @classmethod
    def get_parser(
        cls,
        file_path: Path,
        strategy: ParseStrategy = ParseStrategy.RULE_BASED,
    ) -> BaseParser:
        doc_type = cls.get_doc_type(file_path)
        
        if doc_type == DocumentType.UNKNOWN:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")
        
        parsers_for_type = cls._PARSERS.get(doc_type, {})
        parser_class = parsers_for_type.get(strategy) or parsers_for_type.get(
            ParseStrategy.RULE_BASED
        )
        
        if not parser_class:
            raise ValueError(f"No parser for {doc_type} with strategy {strategy}")
        
        return parser_class()
    
    @classmethod
    def parse(
        cls,
        file_path: Path,
        strategy: ParseStrategy = ParseStrategy.RULE_BASED,
        **kwargs,
    ) -> ParsedDocument:
        """Parse a document with automatic type detection."""
        parser = cls.get_parser(file_path, strategy)
        return parser.parse(file_path, **kwargs)